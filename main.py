# -*- coding: utf-8 -*-

import os
import io
from DrissionPage import ChromiumPage, SessionPage
from DrissionPage.common import Settings
from DrissionPage.errors import ElementNotFoundError
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt
import streamlit as st
import streamlit_ext as ste

Settings.raise_when_ele_not_found = True

version = '2.1.6'


def check_version():
    page = SessionPage()
    # 访问网页
    page.get('https://space.nichx.cn/Version.txt')
    remote_version = page.ele('text:version').text[10:]
    if remote_version == version:
        st.info(f'当前版本为{version} , 是最新版本')
    else:
        st.info(
            f'当前版本为{version} , 最新版本为{remote_version} , 请到 https://share.nichx.cn//s/kaoshibao 下载最新版本')


def download_ques(ID, time):
    page = ChromiumPage()

    url = f'https://www.zaixiankaoshi.com/online/?paperId={ID}'
    page.get(url)

    doc = Document()
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = Pt(11)
    page.wait.eles_loaded('xpath://*[@id="body"]/div[2]/div[1]/div[2]/div[1]/div/div[1]/div/div[1]/div/span[2]')
    number = page.ele('xpath://*[@id="body"]/div[2]/div[1]/div[2]/div[1]/div/div[1]/div/div[1]/div/span[2]').text[2:-1]
    # 打开背题模式
    try:
        button_off = page.s_ele('@@role=switch@@class=el-switch')
        if button_off:
            page.ele('xpath://*[@id="body"]/div[2]/div[1]/div[2]/div[2]/div[2]/div[1]/p[2]/span[2]/div/input').click()
            print('点击背题模式按钮')
            page.wait(0.3, 1.0)
    except ElementNotFoundError:
        print('背题模式已打开')
        page.wait(0.3, 0.9)
    for i in range(int(number)):
        title = f"{i + 1}. {page.ele('@class=qusetion-box').text}"
        doc.add_paragraph(title)
        try:
            ques_img = page.s_ele('@class=qusetion-box').ele('tag:img')
            if ques_img.link:
                ques_img_url = ques_img.attr('src')
                ques_img_url = f'{ques_img_url}'
                page.download(ques_img_url, rf'.\imgs\{ID}\ques', rename=f'ques{i + 1}-title.png')
                page.wait(0.3, 0.6)
            doc.add_picture(rf'.\imgs\{ID}\ques\ques{i + 1}-title.png')
        except ElementNotFoundError:
            pass
        topic = page.ele('@class=topic-type').text
        option = ''
        if topic == '单选题':
            options = page.s_eles('@class^option')
            for j in options:
                try:
                    # 定位选项内的类名以'fr-fic '开头的元素，假设它代表选项图片
                    option_img_url = j.s_ele('tag:img').link
                    # 定位当前选项内的类名以'before-icon'开头的元素
                    x = j.s_ele('@class^before-icon')  # 更改此处，确保使用当前选项的'before-icon'元素
                    # 下载选项图片到指定目录，并重命名
                    page.download(option_img_url, rf'.\imgs\{ID}\option', rename=f'ques{i + 1}-option-{x.text}.png')
                    page.wait(0.3, 0.6)
                    para = doc.add_paragraph()
                    list_j = list(j.text)
                    list_j.insert(1, '.')
                    str_j = ''.join(list_j)
                    run = para.add_run(str_j)  # 添加选项文本
                    img_path = rf'.\imgs\{ID}\option\ques{i + 1}-option-{x.text}.png'
                    run.add_picture(img_path, width=Inches(2.5))
                except Exception as e:
                    list_j = list(j.text)
                    list_j.insert(1, '.')
                    str_j = ''.join(list_j)
                    doc.add_paragraph(str_j)
                    option += str_j + "\n"
            answer = page.ele('@class=right-ans').text.replace('\u2003', ':').rstrip(":")
        elif topic == '判断题':
            options = page.ele('@class^select-left').children('@class^option')
            for j in options:
                list_j = list(j.text)
                list_j.insert(1, '.')
                str_j = ''.join(list_j)
                doc.add_paragraph(str_j)
                option += str_j + "\n"
            answer = page.ele(
                'xpath://*[@id="body"]/div[2]/div[1]/div[2]/div[1]/div/div[3]/div[1]/div/div[1]/div/b').text.replace(
                '\u2003', ':')
        elif topic == '多选题':
            options = page.s_eles('@class^option')
            for j in options:
                try:
                    # 定位选项内的类名以'fr-fic '开头的元素，假设它代表选项图片
                    option_img_url = j.s_ele('tag:img').link
                    # 定位当前选项内的类名以'before-icon'开头的元素
                    x = j.s_ele('@class^before-icon')  # 更改此处，确保使用当前选项的'before-icon'元素
                    # 下载选项图片到指定目录，并重命名
                    page.download(option_img_url, rf'.\imgs\{ID}\option', rename=f'ques{i + 1}-option-{x.text}.png')
                    page.wait(0.3, 0.6)
                    para = doc.add_paragraph()
                    list_j = list(j.text)
                    list_j.insert(1, '.')
                    str_j = ''.join(list_j)
                    run = para.add_run(str_j)  # 添加选项文本
                    img_path = rf'.\imgs\{ID}\option\ques{i + 1}-option-{x.text}.png'
                    run.add_picture(img_path, width=Inches(2.5))
                except Exception as e:
                    list_j = list(j.text)
                    list_j.insert(1, '.')
                    str_j = ''.join(list_j)
                    doc.add_paragraph(str_j)
                    option += str_j + "\n"
            answer = page.ele(
                'xpath://*[@id="body"]/div[2]/div[1]/div[2]/div[1]/div/div[3]/div[1]/div/div[1]/div/b').text.replace(
                '\u2003', ':')
        elif topic == '填空题':
            answer = '正确答案:' + page.s_ele('@class=mt20').text.replace('\u2003', ':')
        elif topic == '简答题':
            answer = '正确答案:' + page.s_ele('@class=mt20').text.replace('\u2003', ':')

        try:
            analysis = page.s_ele('@class^answer-analysis').text.replace('\n', '')
            try:
                analysis_img = page.s_ele('@class^answer-analysis').ele('tag:img')
                if analysis_img.link:
                    analysis_img_url = analysis_img.attr('src')
                    if analysis_img_url == 'https://resource.zaixiankaoshi.com/mini/ai_tag.png':
                        pass
                    else:
                        analysis_img_url = f'{analysis_img_url}'
                        page.download(analysis_img_url, rf'.\imgs\{ID}\analysis', rename=f'ques{i + 1}-analysis.png')
                        page.wait(0.3, 0.6)
            except ElementNotFoundError:
                pass
        except Exception as e:
            print(e)
        try:
            page.ele('@@class:el-button el-button--primary el-button--small@@text():下一题', timeout=5).click()
            page.wait(float(time))
        except Exception as e:
            print(e)
        # 添加答案段落
        doc.add_paragraph(answer)
        doc.add_paragraph(f'解析：{analysis}')
        try:
            doc.add_picture(rf'.\imgs\{ID}\analysis\ques{i + 1}-analysis.png')
        except Exception as e:
            pass
        info = f'第{i + 1}题已完成'
        st.toast(info, icon="✅")
    doc.save(rf'./{ID}.docx')
    doc_stream = io.BytesIO()
    # 将文档保存到BytesIO对象中
    doc.save(doc_stream)
    # 重置缓冲区位置到开始，以便从头读取数据
    doc_stream.seek(0)
    # 返回文档的二进制数据
    return doc_stream.getvalue()


def main():
    st.set_page_config(
        page_title="Kaoshibao Crawler",
        page_icon="🧊",
        initial_sidebar_state="expanded",
        menu_items={
            'Get Help': 'https://space.bilibili.com/34201402',
            'About': "# Kaoshibao Crawler. This is an *extremely* cool app!"
        }
    )
    st.title('Kaoshibao Crawler')
    check_version()
    延迟时间 = st.number_input("Insert a number", step=0.1, value=0.4)
    题库ID = st.text_input("题库ID", None)

    if 题库ID is not None:
        # 假设download_ques现在直接返回一个BytesIO对象或二进制数据
        doc_stream = download_ques(题库ID, 延迟时间)
        # 确保doc_stream可以直接被读取或已转换为BytesIO对象
        if isinstance(doc_stream, io.BytesIO):
            bio = doc_stream  # 如果已经是BytesIO对象，直接使用
        elif isinstance(doc_stream, bytes):
            bio = io.BytesIO(doc_stream)  # 如果是bytes类型，创建一个新的BytesIO对象
        if bio:  # 确保bio有效
            ste.download_button("点击下载docx文件", bio.getvalue(), f"{题库ID}.docx")
        else:
            st.error("无法生成文档，请重试。")


if __name__ == "__main__":
    main()





